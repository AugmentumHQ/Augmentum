"""Document artifact tool — generates PDF and DOCX files from structured content."""

from __future__ import annotations

import io
import json
import os
import re
from typing import TYPE_CHECKING

from augmentum.tools.base import Tool, ToolCategory, ToolResult
from augmentum.utils.logging import get_logger

if TYPE_CHECKING:
    from augmentum.tools.artifact_storage import ArtifactStore

log = get_logger(__name__)


# Patterns that strongly imply the user wants an editable Word doc rather
# than a finalized PDF. Matched case-insensitively on the original user
# request. Order matters only for readability.
_DOCX_INTENT_PATTERNS = (
    r"\b(?:word\s+doc(?:ument)?s?|docx|\.docx|ms\s*word|microsoft\s*word)\b",
    r"\bedit(?:able)?\s+(?:doc|document|version|copy)\b",
    r"\b(?:so\s+i\s+can|that\s+i\s+can)\s+edit\b",
    r"\btrack\s+changes\b",
)
_PDF_INTENT_PATTERNS = (
    r"\bpdf\b|\.pdf\b",
    r"\bprint(?:-?ready|able)\b",
    r"\b(?:final|finalized|print|printable|sign(?:-?off)?)\s+(?:doc|document|report|version)\b",
)

_DOCX_INTENT_RE = re.compile("|".join(_DOCX_INTENT_PATTERNS), re.IGNORECASE)
_PDF_INTENT_RE = re.compile("|".join(_PDF_INTENT_PATTERNS), re.IGNORECASE)


def _sanitize_for_pdf(text: str) -> str:
    """Force ``text`` into Latin-1 so fpdf2's built-in Helvetica can
    render it. Called when ``_setup_pdf_fonts`` couldn't load a real
    Unicode font and the PDF path has to fall back to core fonts.

    Strategy: map common Unicode punctuation to ASCII first (the cases
    that actually look right after conversion), then NFKD-decompose the
    rest and drop anything still outside Latin-1. This avoids hard
    crashes on user content that contains smart quotes, em-dashes,
    ellipses, accented characters, or emoji — which the old code would
    have hit with a bare ``NameError`` because this function was called
    but never defined (see audit finding on PDF font fallback).

    Pair with :func:`content_was_downgraded` to decide whether to
    surface a warning to the user.
    """
    if not text:
        return text
    try:
        text.encode("latin-1")
        return text
    except UnicodeEncodeError:
        pass
    import unicodedata
    # Common smart-punctuation → ASCII (Helvetica renders these fine)
    PUNCT_MAP = {
        "\u2018": "'", "\u2019": "'",  # single quotes
        "\u201c": '"', "\u201d": '"',  # double quotes
        "\u2013": "-", "\u2014": "--",  # dashes
        "\u2026": "...",               # ellipsis
        "\u00a0": " ",                 # nbsp
        "\u2022": "*",                 # bullet
        "\u2010": "-", "\u2011": "-",  # hyphens
    }
    for src, dst in PUNCT_MAP.items():
        text = text.replace(src, dst)
    # Decompose accented characters into base + combining marks so the
    # base letter survives (é → e + combining-acute → e after strip).
    decomposed = unicodedata.normalize("NFKD", text)
    return decomposed.encode("latin-1", "replace").decode("latin-1")


def content_was_downgraded(original: str, sanitized: str) -> bool:
    """True iff ``_sanitize_for_pdf`` actually lost or substituted
    characters — i.e. the PDF won't render exactly what the user
    asked for. Caller should append a warning to ``ToolResult.warnings``.
    """
    return bool(original) and (original != sanitized or "?" in sanitized and "?" not in original)


def infer_document_format(user_text: str) -> str | None:
    """Infer pdf/docx from the user's phrasing.

    Returns ``"docx"`` when the user explicitly asks for a Word /
    editable document, ``"pdf"`` when they ask for a PDF / print-ready
    output, and ``None`` when neither signal is present (caller should
    fall back to its own default — typically ``"pdf"``).

    Used by callers that want to pick a sensible default *before*
    invoking the tool; the tool itself also accepts an explicit
    ``format`` argument which always wins.
    """
    if not user_text:
        return None
    docx = _DOCX_INTENT_RE.search(user_text)
    pdf = _PDF_INTENT_RE.search(user_text)
    if docx and not pdf:
        return "docx"
    if pdf and not docx:
        return "pdf"
    if docx and pdf:
        # Both signals — honour whichever came first in the text, since
        # users often walk back their earlier phrasing ("a PDF — actually
        # send me the editable Word version").
        return "docx" if docx.start() > pdf.start() else "pdf"
    return None


class DocumentTool(Tool):
    """Generate professional documents (PDF or DOCX) from structured sections."""

    def __init__(self, artifact_store: ArtifactStore) -> None:
        self._store = artifact_store

    @property
    def name(self) -> str:
        return "create_document"

    @property
    def description(self) -> str:
        return (
            "Create a professional document (PDF or DOCX) from structured sections. "
            "Each section has a heading, body text (supports markdown), and optional "
            "image_url (e.g. /api/image/abc123) to embed an illustration.\n\n"
            "Choose format='docx' when the user asks for an editable / Word / "
            "docx document, wants to track changes, or says they want to edit "
            "the output. Choose format='pdf' for reports, articles, print-ready "
            "or finalized output, or when unspecified. When in doubt, pdf is "
            "the safer default.\n\n"
            "Returns a download link for the generated file. "
            "Output content directly — do not include meta-commentary, preambles, or placeholder text like [Insert data]."
        )

    @property
    def category(self) -> ToolCategory:
        return ToolCategory.ARTIFACT

    @property
    def input_schema(self) -> dict:
        return {
            "type": "object",
            "properties": {
                "title": {
                    "type": "string",
                    "description": "Document title",
                },
                "format": {
                    "type": "string",
                    "enum": ["pdf", "docx"],
                    "description": (
                        "Output format. Use 'docx' when the user wants an "
                        "editable Word document or plans to keep editing the "
                        "result. Use 'pdf' for reports, articles, print-ready "
                        "output, or when unspecified. Default: pdf."
                    ),
                    "default": "pdf",
                },
                "author": {
                    "type": "string",
                    "description": "Author name (optional)",
                    "default": "",
                },
                "theme": {
                    "type": "string",
                    "description": "Visual theme preset (slate, corporate, modern, emerald, rose). Default: slate",
                    "default": "",
                },
                "sections": {
                    "type": "array",
                    "description": "Ordered list of document sections",
                    "items": {
                        "type": "object",
                        "properties": {
                            "heading": {
                                "type": "string",
                                "description": "Section heading",
                            },
                            "level": {
                                "type": "integer",
                                "description": "Heading level (1-4, default 1)",
                                "default": 1,
                            },
                            "body": {
                                "type": "string",
                                "description": "Section content (markdown supported for PDF)",
                            },
                            "image_url": {
                                "type": "string",
                                "description": (
                                    "URL of an image to embed after this section's text "
                                    "(e.g. /api/image/abc123 from image_generation tool)"
                                ),
                                "default": "",
                            },
                            "image_caption": {
                                "type": "string",
                                "description": "Caption for the embedded image (optional)",
                                "default": "",
                            },
                        },
                        "required": ["heading", "body"],
                    },
                },
            },
            "required": ["title", "sections"],
        }

    async def execute(
        self,
        *,
        title: str = "Document",
        format: str = "pdf",
        author: str = "",
        theme: str = "",
        sections: list | None = None,
        task_id: str = "",
        session_id: str = "",
        **kwargs,
    ) -> ToolResult:
        from augmentum.tools.artifact_normalize import normalize_sections, normalize_str

        title = normalize_str(title, "Document")
        author = normalize_str(author)
        sections = normalize_sections(sections)
        if not sections:
            return ToolResult(success=False, error="No sections provided")

        from augmentum.tools.artifact_sanitize import sanitize_sections
        sections = sanitize_sections(sections)
        if not sections:
            return ToolResult(success=False, error="No content after cleanup")

        fmt = format.lower()
        if fmt not in ("pdf", "docx"):
            return ToolResult(success=False, error=f"Unsupported format: {fmt}")

        # Resolve image paths from URLs before rendering
        resolved_sections = await self._resolve_images(sections)

        # Pre-flight: warn if the PDF path will have to downgrade Unicode
        # content because DejaVu isn't installed on this host.
        warnings: list[str] = []
        if fmt == "pdf" and pdf_render_will_downgrade_unicode(resolved_sections):
            warnings.append(
                "This PDF was rendered with the Helvetica fallback font "
                "because no Unicode TTF is installed on the server — "
                "accented characters were approximated (e.g. é→e) and "
                "smart quotes were replaced with ASCII. Install DejaVu "
                "Sans for full Unicode support."
            )

        try:
            if fmt == "pdf":
                data = _render_pdf(title, author, resolved_sections, theme_name=theme)
            else:
                data = _render_docx(title, author, resolved_sections, theme_name=theme)

            safe_title = re.sub(r'[^\w\s-]', '', title).strip().replace(' ', '_')[:60]
            filename = f"{safe_title}.{fmt}"

            image_count = sum(
                1 for s in resolved_sections if s.get("_image_path")
            )

            info = await self._store.save(
                data=data,
                filename=filename,
                fmt=fmt,
                task_id=task_id,
                session_id=session_id,
                display_name=f"{title}.{fmt}",
                user_id=Tool.extract_user_id(kwargs),
                metadata={
                    "page_type": "document",
                    "section_count": len(sections),
                    "image_count": image_count,
                    "author": author,
                },
                source_json=json.dumps({
                    "type": "document",
                    "title": title,
                    "author": author,
                    "format": fmt,
                    "theme": theme,
                    "sections": sections,
                }),
            )

            from augmentum.tools.base import make_artifact_card

            img_str = f", {image_count} image{'s' if image_count != 1 else ''}" if image_count else ""
            summary = (
                f"{fmt.upper()} document '{title}' is ready — "
                f"{len(sections)} section{'s' if len(sections) != 1 else ''}{img_str}. "
                "Available in the artifact library."
            )
            card = make_artifact_card(
                info,
                kind="artifact",
                title=title,
                subtitle=f"by {author}" if author else fmt.upper(),
                summary=summary,
                preview={
                    "artifact_kind": "document",
                    "format": fmt,
                    "size_bytes": info.get("size_bytes", 0),
                    "sections": [
                        {"heading": s.get("heading", f"Section {i+1}")}
                        for i, s in enumerate(sections)
                    ],
                    "image_count": image_count,
                },
            )
            from augmentum.tools.base import format_output_with_warnings

            return ToolResult(
                success=True,
                output=format_output_with_warnings(summary, warnings),
                metadata=info,
                warnings=warnings,
                card=card,
            )
        except Exception as e:
            log.error("document_creation_failed", error=str(e), exc_info=True)
            return ToolResult(success=False, error=f"Document creation failed: {e}")

    async def _resolve_images(self, sections: list) -> list:
        """Resolve image_url references to filesystem paths.

        Supports:
        - /api/image/<image_id> — looks up via ImagePersistence
        - /api/artifacts/<artifact_id>/download — looks up via ArtifactStore
        """
        resolved = []
        for section in sections:
            s = dict(section)
            image_url = s.get("image_url", "")
            if image_url:
                path = await self._resolve_image_path(image_url)
                if path:
                    s["_image_path"] = str(path)
                else:
                    log.warning("image_resolve_failed", url=image_url)
            resolved.append(s)
        return resolved

    async def _resolve_image_path(self, url: str) -> str | None:
        """Resolve an image URL to a local filesystem path."""
        # /api/image/<image_id>
        m = re.match(r"/api/image/([a-zA-Z0-9_-]+)", url)
        if m:
            image_id = m.group(1)
            return await self._resolve_image_id(image_id)

        # /api/artifacts/<artifact_id>/download
        m = re.match(r"/api/artifacts/([a-zA-Z0-9_-]+)/download", url)
        if m:
            artifact_id = m.group(1)
            info = await self._store.get(artifact_id)
            if info:
                file_path = self._store.get_file_path(info["path"])
                if file_path:
                    return str(file_path)

        return None

    async def _resolve_image_id(self, image_id: str) -> str | None:
        """Look up an image_id via the ImagePersistence store."""
        try:
            # Access image persistence from the same DB connection
            db = self._store._db
            async with db.execute(
                "SELECT file_path FROM image_generations WHERE image_id = ?",
                (image_id,),
            ) as cursor:
                row = await cursor.fetchone()
                if row and row[0] and os.path.exists(row[0]):
                    return row[0]
        except Exception as e:
            log.debug("image_id_lookup_failed", image_id=image_id, error=str(e))
        return None


# ---------------------------------------------------------------------------
# PDF rendering via fpdf2
# ---------------------------------------------------------------------------

_TOC_TITLE_SIZE = 16
_TOC_ENTRY_SIZE = 10
_TOC_LINE_HEIGHT = 7
_TOC_INDENT = 10  # extra indent per heading level beyond 1
_CAPTION_SIZE = 8


def _md_to_html(text: str) -> str:
    """Convert markdown to HTML using markdown-it-py.

    Falls back to basic manual conversion if the library isn't available.
    """
    try:
        from markdown_it import MarkdownIt
        md = MarkdownIt("commonmark", {"typographer": True})
        return md.render(text)
    except ImportError:
        # Fallback: basic markdown → HTML for bold, italic, lists
        html = text
        html = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", html)
        html = re.sub(r"\*(.+?)\*", r"<i>\1</i>", html)
        html = re.sub(r"^- (.+)$", r"<li>\1</li>", html, flags=re.MULTILINE)
        html = re.sub(r"(<li>.*</li>)", r"<ul>\1</ul>", html, flags=re.DOTALL)
        html = html.replace("\n\n", "<br><br>")
        return html


# Ordered preference of font families for the PDF body face. Each entry maps
# the four fpdf2 styles to a TTF path; missing italic/bolditalic fall back to
# regular/bold. Liberation Sans (Arial-metric, and the only system family that
# ships TRUE italic + bold-italic faces) is preferred for a clean, professional
# look — DejaVu has no installed italics so emphasis was being faked from the
# regular weight. DejaVu Sans stays as the broad-Unicode fallback; the Windows
# entry covers bare dev hosts. Helvetica core is the last resort (the caller
# falls back to it + ASCII sanitization when this returns None).
_PDF_FONT_FAMILIES: tuple[dict[str, str], ...] = (
    {  # Liberation Sans — Debian/Ubuntu (ships with Chrome's deps)
        "regular": "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
        "bold": "/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf",
        "italic": "/usr/share/fonts/truetype/liberation/LiberationSans-Italic.ttf",
        "bolditalic": "/usr/share/fonts/truetype/liberation/LiberationSans-BoldItalic.ttf",
    },
    {  # Liberation Sans — Arch / alt layout
        "regular": "/usr/share/fonts/liberation/LiberationSans-Regular.ttf",
        "bold": "/usr/share/fonts/liberation/LiberationSans-Bold.ttf",
        "italic": "/usr/share/fonts/liberation/LiberationSans-Italic.ttf",
        "bolditalic": "/usr/share/fonts/liberation/LiberationSans-BoldItalic.ttf",
    },
    {  # DejaVu Sans — broad Unicode coverage (no installed italics → faked)
        "regular": "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "bold": "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
    },
    {  # DejaVu Sans — Arch
        "regular": "/usr/share/fonts/TTF/DejaVuSans.ttf",
        "bold": "/usr/share/fonts/TTF/DejaVuSans-Bold.ttf",
    },
    {  # Windows dev host — Arial
        "regular": "C:/Windows/Fonts/arial.ttf",
        "bold": "C:/Windows/Fonts/arialbd.ttf",
        "italic": "C:/Windows/Fonts/ariali.ttf",
        "bolditalic": "C:/Windows/Fonts/arialbi.ttf",
    },
)

# Back-compat flat path list (some call sites / tests referenced this name).
_PDF_UNICODE_FONT_PATHS = tuple(
    p for fam in _PDF_FONT_FAMILIES for p in fam.values()
)


def _select_pdf_font_family() -> dict[str, str] | None:
    """Return the first font family whose regular face exists on disk."""
    for fam in _PDF_FONT_FAMILIES:
        if os.path.exists(fam.get("regular", "")):
            return fam
    return None


def _unicode_fonts_available() -> bool:
    """True iff ``_setup_pdf_fonts`` would find a Unicode TTF family."""
    return _select_pdf_font_family() is not None


def pdf_render_will_downgrade_unicode(sections: list) -> bool:
    """Pre-flight check: will the PDF path have to substitute/drop
    characters from this content because DejaVu Sans isn't installed?

    Returns True only when BOTH conditions hold:
    1. No Unicode-capable TTF is discoverable on disk, AND
    2. At least one section contains characters outside Latin-1.

    When True, callers should append a user-visible warning — otherwise
    the user silently gets a PDF with mojibake / dropped glyphs, which
    was the HIGH-severity finding in the artifact pipeline audit.
    """
    if _unicode_fonts_available():
        return False
    for section in sections or []:
        if not isinstance(section, dict):
            continue
        for key in ("heading", "body", "image_caption"):
            text = section.get(key, "")
            if not isinstance(text, str) or not text:
                continue
            try:
                text.encode("latin-1")
            except UnicodeEncodeError:
                return True
    return False


def _setup_pdf_fonts(pdf) -> bool:
    """Load the best available Unicode TTF family as "Doc" with all four
    styles. Returns True on success; the caller falls back to Helvetica +
    ASCII sanitization when this returns False.

    Liberation Sans (preferred) ships real italic + bold-italic faces, so
    emphasis renders properly instead of being faked from the regular weight.
    Families without those faces fall regular/bold in for I/BI.
    """
    fam = _select_pdf_font_family()
    if not fam:
        return False

    regular = fam["regular"]
    bold = fam["bold"] if os.path.exists(fam.get("bold", "")) else regular
    italic = fam["italic"] if os.path.exists(fam.get("italic", "")) else regular
    bolditalic = (
        fam["bolditalic"] if os.path.exists(fam.get("bolditalic", "")) else bold
    )

    try:
        pdf.add_font("Doc", "", regular)
        pdf.add_font("Doc", "B", bold)
        pdf.add_font("Doc", "I", italic)
        pdf.add_font("Doc", "BI", bolditalic)
        return True
    except Exception as e:
        log.debug("ttf_font_load_failed", error=str(e))
        return False


def _render_pdf(
    title: str,
    author: str,
    sections: list,
    theme_name: str = "",
    design: dict | None = None,
) -> bytes:
    """Render sections to a professional-looking PDF using fpdf2 + theme system.

    `design` (optional) applies font_size_scale + line_height + density +
    accent_override on top of the theme. None = use theme defaults.
    """
    from fpdf import FPDF

    from augmentum.tools.artifact_theme import FONT_FAMILY_PDF, apply_design, get_theme

    theme = apply_design(get_theme(theme_name), design)

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=theme.margin_bottom)
    pdf.set_margins(theme.margin_left, theme.margin_top, theme.margin_right)

    # Try Unicode fonts first; fall back to Helvetica with sanitization.
    # design.font_family picks the fallback face (Times for serif, Courier
    # for mono, Helvetica otherwise). Unicode fonts only ship as sans — if
    # we loaded them, design.font_family is honored only for non-Unicode
    # builds. Pro-tool tradeoff: glyph coverage > font-family fidelity.
    has_unicode = _setup_pdf_fonts(pdf)
    family_pref = (design or {}).get("font_family", "system")
    fallback = FONT_FAMILY_PDF.get(family_pref, "Helvetica")
    font_family = "Doc" if has_unicode else fallback

    if not has_unicode:
        title = _sanitize_for_pdf(title)
        author = _sanitize_for_pdf(author)
        for section in sections:
            for key in ("heading", "body", "image_caption"):
                if key in section and isinstance(section[key], str):
                    section[key] = _sanitize_for_pdf(section[key])

    pdf.set_title(title)
    if author:
        pdf.set_author(author)
    pdf.set_creator("Augmentum")

    ff = font_family  # shorthand

    # -- Cover page --
    from datetime import datetime

    pdf.add_page()
    # Absolute-position the cover; never let a long title trip a page break.
    pdf.set_auto_page_break(auto=False)

    # Full-bleed accent header block with the title reversed in white.
    band_h = max(64.0, theme.cover_bar_height * 3.4)
    bar_r, bar_g, bar_b = theme.rgb(theme.accent_dark)
    pdf.set_fill_color(bar_r, bar_g, bar_b)
    pdf.rect(0, 0, pdf.w, band_h, "F")
    # Brighter accent stripe along the bottom edge of the band.
    ar, ag, ab = theme.rgb(theme.accent)
    pdf.set_fill_color(ar, ag, ab)
    pdf.rect(0, band_h, pdf.w, 2.5, "F")

    inner_w = pdf.w - theme.margin_left - theme.margin_right

    # Kicker (author or "Document"), uppercase, inside the band.
    pdf.set_xy(theme.margin_left, band_h * 0.30)
    pdf.set_text_color(255, 255, 255)
    pdf.set_font(ff, "B", 10)
    pdf.cell(0, 6, (author or "Document").upper(), new_x="LMARGIN", new_y="NEXT")

    # Title, large and bold, reversed in white.
    pdf.set_xy(theme.margin_left, band_h * 0.30 + 9)
    pdf.set_font(ff, "B", theme.title_size)
    pdf.multi_cell(inner_w, theme.title_size * 0.52, title, align="L")

    # Below the band: author + date stacked in slate.
    pdf.set_xy(theme.margin_left, band_h + 18)
    if author:
        pdf.set_text_color(*theme.rgb(theme.text))
        pdf.set_font(ff, "", 13)
        pdf.cell(0, 8, author, new_x="LMARGIN", new_y="NEXT")
        pdf.set_x(theme.margin_left)
    pdf.set_text_color(*theme.rgb(theme.text_muted))
    pdf.set_font(ff, "", 10)
    pdf.cell(0, 7, datetime.now().strftime("%B %d, %Y"), new_x="LMARGIN", new_y="NEXT")

    # Credit line + hairline pinned near the foot of the cover.
    pdf.set_xy(theme.margin_left, pdf.h - theme.margin_bottom - 8)
    pdf.set_draw_color(*theme.rgb(theme.border))
    pdf.set_line_width(0.3)
    pdf.line(theme.margin_left, pdf.get_y(), pdf.w - theme.margin_right, pdf.get_y())
    pdf.ln(2)
    pdf.set_x(theme.margin_left)
    pdf.set_text_color(*theme.rgb(theme.text_muted))
    pdf.set_font(ff, "I", 8)
    pdf.cell(0, 6, "Generated by Augmentum", new_x="LMARGIN", new_y="NEXT")

    # Restore normal flow for the content pages.
    pdf.set_auto_page_break(auto=True, margin=theme.margin_bottom)

    # -- TOC placeholder page (only for documents with 4+ sections) --
    has_toc = len(sections) >= 4
    toc_page = 0
    if has_toc:
        pdf.add_page()
        toc_page = pdf.page

    # -- Content pages --
    pdf.add_page()

    content_width = pdf.w - 50  # usable width between margins

    # Collect heading entries for the TOC
    toc_entries: list[tuple[str, int, int, int]] = []

    for section in sections:
        heading = section.get("heading", "")
        level = section.get("level", 1)
        body = section.get("body", "")
        image_path = section.get("_image_path", "")
        image_caption = section.get("image_caption", "")

        # Models sometimes return body as a list instead of a string
        if isinstance(body, list):
            body = "\n".join(str(item) for item in body)

        # Heading
        if heading:
            size, style = {
                1: (theme.h1_size, "B"),
                2: (theme.h2_size, "B"),
                3: (theme.h3_size, "B"),
                4: (theme.h4_size, "BI"),
            }.get(level, (theme.h1_size, "B"))

            pdf.ln(4 if level > 1 else 9)

            link_dest = pdf.add_link()
            pdf.set_link(link_dest)
            toc_entries.append((heading, level, pdf.page, link_dest))

            if level == 1:
                # Accent tab bar to the left of the heading, heading in accent,
                # then a short accent underline — a designed section marker
                # instead of a full-width hairline.
                ar2, ag2, ab2 = theme.rgb(theme.accent)
                tab_y = pdf.get_y()
                pdf.set_fill_color(ar2, ag2, ab2)
                pdf.rect(theme.margin_left, tab_y + 1.5, 2.2, size * 0.42, "F")
                pdf.set_text_color(ar2, ag2, ab2)
                pdf.set_font(ff, style, size)
                pdf.set_x(theme.margin_left + 6)
                pdf.multi_cell(0, size * 0.6, heading, new_x="LMARGIN", new_y="NEXT")
                y = pdf.get_y() + 0.5
                pdf.set_draw_color(ar2, ag2, ab2)
                pdf.set_line_width(0.8)
                pdf.line(theme.margin_left, y, theme.margin_left + 26, y)
                pdf.ln(3.5)
            else:
                r, g, b = theme.rgb(theme.text)
                pdf.set_text_color(r, g, b)
                pdf.set_font(ff, style, size)
                pdf.multi_cell(0, size * 0.6, heading, new_x="LMARGIN", new_y="NEXT")

        # Body — markdown → HTML → write_html for rich rendering
        if body:
            r, g, b = theme.rgb(theme.text)
            pdf.set_text_color(r, g, b)
            pdf.set_font(ff, "", theme.body_size)
            pdf.ln(2)

            try:
                html = _md_to_html(body)
                # write_html handles paragraphs, bold, italic, lists, tables, code.
                # li_prefix_color overrides fpdf2's default dark-red (190,0,0)
                # bullet/number marker color with the theme accent.
                try:
                    pdf.write_html(html, li_prefix_color=theme.rgb(theme.accent))
                except TypeError:
                    # Older fpdf2 without the li_prefix_color kwarg.
                    pdf.write_html(html)
            except Exception:
                # Fallback: plain text rendering if HTML pipeline fails
                log.debug("write_html_fallback", exc_info=True)
                for paragraph in body.split("\n\n"):
                    paragraph = paragraph.strip()
                    if not paragraph:
                        continue
                    pdf.multi_cell(
                        0, theme.line_height, paragraph,
                        new_x="LMARGIN", new_y="NEXT",
                        markdown=True,
                    )
                    pdf.ln(2)

        # Embed image after section text
        if image_path and os.path.exists(image_path):
            _embed_pdf_image(pdf, image_path, image_caption, content_width, ff)

    # -- Fill in the TOC page --
    if has_toc and toc_entries:
        _render_toc(pdf, toc_page, toc_entries, ff, theme)

    # Running footer with page numbers and optional title
    page_count = pdf.pages_count
    mr, mg, mb = theme.rgb(theme.text_muted)
    br, bg, bb = theme.rgb(theme.border)
    for i in range(1, page_count + 1):
        pdf.page = i
        # Skip cover page
        if i == 1:
            continue
        pdf.set_y(-15)
        # Thin rule above footer
        pdf.set_draw_color(br, bg, bb)
        pdf.set_line_width(0.2)
        pdf.line(theme.margin_left, pdf.get_y() - 2,
                 pdf.w - theme.margin_right, pdf.get_y() - 2)
        pdf.set_text_color(mr, mg, mb)
        pdf.set_font(ff, "", 8)
        # Title on left (muted), page number on right (accent, bold)
        pdf.cell(0, 10, title[:60], align="L")
        pdf.set_x(theme.margin_left)
        pdf.set_text_color(*theme.rgb(theme.accent))
        pdf.set_font(ff, "B", 8)
        pdf.cell(0, 10, f"{i} / {page_count}", align="R")

    return bytes(pdf.output())


def _render_toc(
    pdf,
    toc_page: int,
    entries: list[tuple[str, int, int, int]],
    ff: str = "Helvetica",
    theme=None,
) -> None:
    """Switch to the TOC page and render a clickable table of contents.

    Each entry is (heading_text, level, page_number, link_dest). Rendered as
    clean ruled rows (title left, page number right) — no dotted leaders.
    """
    # Theme colors, with neutral fallbacks if no theme was passed.
    accent = theme.rgb(theme.accent) if theme else (37, 99, 235)
    text_col = theme.rgb(theme.text) if theme else (15, 23, 42)
    muted = theme.rgb(theme.text_muted) if theme else (100, 116, 139)
    border = theme.rgb(theme.border) if theme else (226, 232, 240)

    pdf.page = toc_page
    left_margin = pdf.l_margin
    right_margin = pdf.r_margin
    usable_width = pdf.w - left_margin - right_margin

    # TOC title in accent, with a short underline.
    pdf.set_y(28)
    pdf.set_x(left_margin)
    pdf.set_font(ff, "B", 22)
    pdf.set_text_color(*accent)
    pdf.cell(0, 12, "Contents", new_x="LMARGIN", new_y="NEXT")
    y = pdf.get_y() + 1
    pdf.set_draw_color(*accent)
    pdf.set_line_width(1.2)
    pdf.line(left_margin, y, left_margin + 30, y)
    pdf.ln(9)

    for heading_text, level, page_num, link_dest in entries:
        indent = _TOC_INDENT * (level - 1)
        x_start = left_margin + indent
        page_str = str(page_num)
        page_w = pdf.get_string_width(page_str) + 3
        title_w = usable_width - indent - page_w
        row_h = 9

        if level == 1:
            pdf.set_font(ff, "B", 11)
            pdf.set_text_color(*text_col)
        else:
            pdf.set_font(ff, "", 10)
            pdf.set_text_color(*muted)

        # Truncate over-long headings to keep the page number aligned.
        ht = heading_text
        if pdf.get_string_width(ht) > title_w:
            while pdf.get_string_width(ht + "...") > title_w and len(ht) > 8:
                ht = ht[:-1]
            ht = ht.rstrip() + "..."

        pdf.set_x(x_start)
        pdf.cell(title_w, row_h, ht, link=link_dest)
        pdf.set_x(pdf.w - right_margin - page_w)
        pdf.set_text_color(*muted)
        pdf.cell(
            page_w, row_h, page_str,
            align="R", new_x="LMARGIN", new_y="NEXT", link=link_dest,
        )

        # Hairline divider beneath the row.
        yb = pdf.get_y() + 0.5
        pdf.set_draw_color(*border)
        pdf.set_line_width(0.2)
        pdf.line(left_margin, yb, pdf.w - right_margin, yb)
        pdf.ln(2.5)


def _embed_pdf_image(
    pdf, image_path: str, caption: str, max_width: float,
    ff: str = "Helvetica",
) -> None:
    """Embed an image in the PDF, centered with optional caption."""
    try:
        from PIL import Image as PILImage

        with PILImage.open(image_path) as img:
            img_w, img_h = img.size

        # Scale to fit within content width, max 60% page height
        max_height = (pdf.h - 40) * 0.6
        scale_w = max_width / img_w
        scale_h = max_height / img_h
        scale = min(scale_w, scale_h, 1.0)  # never upscale

        display_w = img_w * scale
        display_h = img_h * scale

        # Check if we need a new page
        remaining = pdf.h - pdf.get_y() - 20  # bottom margin
        needed = display_h + (10 if caption else 5)
        if needed > remaining:
            pdf.add_page()

        # Center the image
        x_offset = (pdf.w - display_w) / 2
        pdf.ln(4)
        pdf.image(image_path, x=x_offset, w=display_w, h=display_h)

        # Caption
        if caption:
            pdf.ln(2)
            pdf.set_font(ff, "I", _CAPTION_SIZE)
            pdf.set_text_color(100, 100, 100)
            pdf.multi_cell(0, 4, caption, align="C", new_x="LMARGIN", new_y="NEXT")
            pdf.set_text_color(0, 0, 0)

        pdf.ln(4)
    except Exception as e:
        log.warning("pdf_image_embed_failed", path=image_path, error=str(e))
        # Fall back to a placeholder note
        pdf.set_font(ff, "I", _CAPTION_SIZE)
        pdf.set_text_color(150, 150, 150)
        pdf.multi_cell(
            0, _TOC_LINE_HEIGHT, f"[Image: {caption or image_path}]",
            align="C", new_x="LMARGIN", new_y="NEXT",
        )
        pdf.set_text_color(0, 0, 0)


# ---------------------------------------------------------------------------
# DOCX rendering via python-docx
# ---------------------------------------------------------------------------


_CODE_FENCE_RE = re.compile(r"^```(\w*)\s*$")
_BULLET_LINE_RE = re.compile(r"^(?P<indent>\s*)(?P<marker>[-*•])\s+(?P<text>.*)$")
_NUMBERED_LINE_RE = re.compile(r"^(?P<indent>\s*)\d+[.\)]\s+(?P<text>.*)$")
_INLINE_HEADING_RE = re.compile(r"^(#{1,4})\s+(.*)$")


def _render_docx_body(doc, body: str) -> None:
    """Render a markdown-flavoured section body into a python-docx Document.

    Covers the common cases the DOCX path previously dropped silently
    (code blocks, inline code, nested lists, inline headings) so the
    same ``body`` text produces comparable output in PDF and DOCX —
    closing the parity gap flagged as CRITICAL in the artifact audit.
    Tables and hyperlinks are not yet implemented; they render as plain
    text and are tracked as a follow-up.
    """
    from docx.shared import Pt, RGBColor

    lines = body.split("\n")
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # Fenced code block ``` … ```
        fence = _CODE_FENCE_RE.match(stripped)
        if fence:
            code_lines: list[str] = []
            i += 1
            while i < n and not _CODE_FENCE_RE.match(lines[i].strip()):
                code_lines.append(lines[i])
                i += 1
            # Consume the closing fence (or EOF)
            if i < n:
                i += 1
            code_text = "\n".join(code_lines)
            p = doc.add_paragraph()
            run = p.add_run(code_text)
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(40, 40, 60)
            continue

        # Inline heading (## Section Inside Body)
        heading_m = _INLINE_HEADING_RE.match(stripped)
        if heading_m:
            level = min(len(heading_m.group(1)) + 1, 4)
            doc.add_heading(heading_m.group(2), level=level)
            i += 1
            continue

        # Bullet list (with optional nesting — every 2 leading spaces /
        # every tab is one level deeper, so "  - sub" becomes level 2).
        bullet_m = _BULLET_LINE_RE.match(line)
        if bullet_m:
            while i < n:
                m = _BULLET_LINE_RE.match(lines[i])
                if not m:
                    break
                indent = m.group("indent").replace("\t", "  ")
                level = min(max(len(indent) // 2, 0), 3)
                style = "List Bullet" if level == 0 else f"List Bullet {level + 1}"
                p = doc.add_paragraph(style=style)
                _add_formatted_text(p, m.group("text").strip())
                i += 1
            continue

        # Numbered list (same nesting logic)
        num_m = _NUMBERED_LINE_RE.match(line)
        if num_m:
            while i < n:
                m = _NUMBERED_LINE_RE.match(lines[i])
                if not m:
                    break
                indent = m.group("indent").replace("\t", "  ")
                level = min(max(len(indent) // 2, 0), 3)
                style = "List Number" if level == 0 else f"List Number {level + 1}"
                p = doc.add_paragraph(style=style)
                _add_formatted_text(p, m.group("text").strip())
                i += 1
            continue

        # Blank line → paragraph break (no-op; add_paragraph without text)
        if not stripped:
            i += 1
            continue

        # Regular paragraph — collect consecutive non-blank, non-list,
        # non-code lines and join them with spaces (markdown paragraph
        # semantics: soft-wrap within a paragraph).
        para_lines = [line]
        i += 1
        while i < n:
            nxt = lines[i]
            nxt_stripped = nxt.strip()
            if not nxt_stripped:
                break
            if (_CODE_FENCE_RE.match(nxt_stripped)
                    or _BULLET_LINE_RE.match(nxt)
                    or _NUMBERED_LINE_RE.match(nxt)
                    or _INLINE_HEADING_RE.match(nxt_stripped)):
                break
            para_lines.append(nxt)
            i += 1
        p = doc.add_paragraph()
        _add_formatted_text(p, " ".join(pl.strip() for pl in para_lines))


def _render_docx(
    title: str,
    author: str,
    sections: list,
    theme_name: str = "",
    design: dict | None = None,
) -> bytes:
    """Render sections to a DOCX file using python-docx.

    `design` (optional) applies font_family + font_size_scale to the base
    style; line_height and density are no-ops here (DOCX uses paragraph
    spacing primitives that don't map cleanly to the design enum).
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    from augmentum.tools.artifact_theme import FONT_FAMILY_DOCX, apply_design, get_theme

    theme = apply_design(get_theme(theme_name), design)
    family_pref = (design or {}).get("font_family", "system")
    base_font = FONT_FAMILY_DOCX.get(family_pref, "Calibri")

    doc = Document()

    # Core properties
    doc.core_properties.title = title
    if author:
        doc.core_properties.author = author

    # Default font — scales with theme.body_size (which design has already
    # multiplied by font_size_scale via apply_design).
    style = doc.styles["Normal"]
    font = style.font
    font.name = base_font
    font.size = Pt(round(theme.body_size + 1))   # body_size baseline is 10; +1 ≈ Word's 11pt

    # Title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if author:
        author_para = doc.add_paragraph()
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = author_para.add_run(author)
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(100, 100, 100)

    # Subtle "Generated by" line
    gen_para = doc.add_paragraph()
    gen_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = gen_para.add_run("Generated by Augmentum")
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(150, 150, 150)

    doc.add_page_break()

    # Sections
    for section in sections:
        heading = section.get("heading", "")
        level = min(section.get("level", 1), 4)
        body = section.get("body", "")
        image_path = section.get("_image_path", "")
        image_caption = section.get("image_caption", "")

        if heading:
            doc.add_heading(heading, level=level)

        if body:
            _render_docx_body(doc, body)

        # Embed image after section text
        if image_path and os.path.exists(image_path):
            _embed_docx_image(doc, image_path, image_caption)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _embed_docx_image(doc, image_path: str, caption: str) -> None:
    """Embed an image in the DOCX, centered with optional caption."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    try:
        from PIL import Image as PILImage

        with PILImage.open(image_path) as img:
            img_w, img_h = img.size

        # Scale to max 5.5 inches wide (standard page width minus margins).
        # Portrait images get a smaller cap so tall photos don't dominate
        # the page; landscape images use the full content width.
        max_width_inches = 5.5
        if img_w > img_h:
            width = min(max_width_inches, img_w / 96)  # assume 96 DPI
        else:
            width = min(max_width_inches * 0.7, img_w / 96)

        width = min(width, max_width_inches)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(image_path, width=Inches(width))

        if caption:
            cap_p = doc.add_paragraph()
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cap_p.add_run(caption)
            run.font.size = Pt(9)
            run.font.italic = True
            run.font.color.rgb = RGBColor(100, 100, 100)
    except Exception as e:
        log.warning("docx_image_embed_failed", path=image_path, error=str(e))
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"[Image: {caption or image_path}]")
        run.font.italic = True
        run.font.color.rgb = RGBColor(150, 150, 150)


_INLINE_MD_RE = re.compile(
    r"(\*\*[^*]+\*\*"         # **bold**
    r"|\*[^*]+\*"             # *italic*
    r"|`[^`]+`)",             # `inline code`
)


def _add_formatted_text(paragraph, text: str) -> None:
    """Parse basic inline markdown (bold / italic / inline code) and
    append formatted runs to an existing python-docx paragraph."""
    from docx.shared import Pt
    for part in _INLINE_MD_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) >= 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`") and len(part) >= 2:
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        else:
            paragraph.add_run(part)
